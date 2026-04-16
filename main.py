import os
import logging
import base64
import io
import json
import random
import requests
import urllib3
import time
import re
from PIL import Image
import telebot
from telebot import apihelper
from docx import Document

apihelper.SESSION = requests.Session()
apihelper.SESSION.verify = False

from matrix_parser import extract_paragraph_full_text, format_matrix

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

BOT_TOKEN = ""
OPENROUTER_API_KEY = ""

MODEL = "google/gemini-2.5-flash"
URL = "https://openrouter.ai/api/v1/chat/completions"

MAX_CONTEXT_CHARS = 6000
MAX_SEARCH_BLOCKS = 5

bot = telebot.TeleBot(BOT_TOKEN)
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

MEMORY_FILE = "memory.json"
user_conversations = {}
user_current_task = {}  # Сохраняем текущую задачу для каждого пользователя


def load_tasks():
    """Загружает задачи из JSON файла"""
    try:
        with open("tasks.json", "r", encoding="utf-8") as f:
            data = json.load(f)
            return data.get("tasks", [])
    except FileNotFoundError:
        print("Файл tasks.json не найден!")
        return []
    except json.JSONDecodeError as e:
        print(f"Ошибка парсинга JSON: {e}")
        return []


TASKS = load_tasks()


def get_task_by_id(task_id):
    """Получить задачу по ID"""
    for task in TASKS:
        if task["id"] == task_id:
            return task
    return None


def get_tasks_by_variant(variant_num):
    """Получить все задачи варианта"""
    return [task for task in TASKS if task["variant"] == variant_num]


def get_random_task_from_json(task_type=None):
    if not TASKS:
        return None

    if task_type:
        filtered = [task for task in TASKS if task.get("type") == task_type]
        if filtered:
            return random.choice(filtered)

    return random.choice(TASKS)


def get_task_by_variant_and_number(variant, task_number):
    for task in TASKS:
        if task["variant"] == variant and task["task_number"] == task_number:
            return task
    return None


def format_task_from_json(task):
    if not task:
        return None

    text = f"<b> {task['title']}</b>\n\n"
    text += f"<b>📝 Задание:</b>\n{task['question']}\n\n"
    text += "<i>Реши и пришли своё решение — я проверю.\nИли напиши «покажи ответ» или «подскажи».</i>"

    return text


def format_answer_from_json(task):
    """Форматирует ответ из JSON для Telegram"""
    if not task:
        return None

    text = f"<b> {task['title']} - Ответ</b>\n\n"
    text += f"<b>✅ Решение:</b>\n<pre>{task['answer']}</pre>"

    return text


def format_hint_from_json(task):
    """Форматирует подсказку из JSON для Telegram (первые 200 символов ответа)"""
    if not task:
        return None

    hint = task['answer'][:300] + "..." if len(task['answer']) > 300 else task['answer']

    text = f"<b>📘 {task['title']} - Подсказка</b>\n\n"
    text += f"<b>💡 Подсказка:</b>\n<pre>{hint}</pre>\n\n"
    text += "<i>Это только часть решения. Напиши «покажи ответ» для полного решения.</i>"

    return text


def load_memory():
    global user_conversations, user_current_task
    try:
        with open(MEMORY_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
            user_conversations = data.get("conversations", {})
            user_current_task = data.get("current_tasks", {})
    except:
        user_conversations = {}
        user_current_task = {}


def save_memory():
    with open(MEMORY_FILE, "w", encoding="utf-8") as f:
        data = {
            "conversations": user_conversations,
            "current_tasks": user_current_task
        }
        json.dump(data, f, ensure_ascii=False, indent=2)


def load_docx_text():
    try:
        doc = Document("text.docx")
        lines = []
        for p in doc.paragraphs:
            text = extract_paragraph_full_text(p).strip()
            if text:
                lines.append(text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    lines.append(" | ".join(row_text))
        return "\n".join(lines)
    except FileNotFoundError:
        print("Файл text.docx не найден!")
        return ""
    except Exception as e:
        print(f"Ошибка загрузки docx: {e}")
        return ""


def parse_book_from_docx():
    try:
        doc = Document("text.docx")
    except Exception as e:
        print(f"Ошибка загрузки docx: {e}")
        return {"full_text": TEXT, "sections": [], "tasks": [], "examples": [], "answers": ""}

    paragraphs = doc.paragraphs
    sections = []
    tasks = []
    examples = []
    answers = ""

    current_section = None
    section_content_lines = []
    in_answers = False
    in_exercises = False

    def flush_section():
        nonlocal section_content_lines
        if current_section and section_content_lines:
            sections.append({
                "title": current_section,
                "content": "\n".join(section_content_lines)
            })
        section_content_lines = []

    def is_top_level_list(p):
        if p.style.name != "List Paragraph":
            return False
        pPr = p._element.pPr
        if pPr is None or pPr.numPr is None:
            return False
        ilvl = pPr.numPr.ilvl
        return ilvl is None or ilvl.val == 0

    i = 0
    while i < len(paragraphs):
        p = paragraphs[i]
        text = extract_paragraph_full_text(p).strip()

        if p.style.name == "Heading 1" and p.text.strip():
            flush_section()
            current_section = p.text.strip()
            in_exercises = False
            in_answers = "ответ" in current_section.lower()
            i += 1
            continue

        if in_answers:
            if text:
                answers += text + "\n"
            i += 1
            continue

        if p.text.strip() == "Упражнения":
            in_exercises = True
            section_content_lines.append("\n[Упражнения]")
            i += 1
            continue

        if in_exercises and text:
            section_content_lines.append(text)

            if is_top_level_list(p):
                task_lines = [text]
                j = i + 1
                while j < len(paragraphs) and j < i + 8:
                    np_ = paragraphs[j]
                    nt = extract_paragraph_full_text(np_).strip()
                    if not nt:
                        j += 1
                        continue
                    if is_top_level_list(np_) or np_.style.name == "Heading 1":
                        break
                    if np_.style.name == "List Paragraph" and nt:
                        task_lines.append(nt)
                    j += 1

                full_task = "\n".join(task_lines)
                if len(full_task.strip()) > 15:
                    label = current_section or "Упражнения"
                    tasks.append(f"[{label}]\n{full_task[:600]}")

            i += 1
            continue

        if text:
            section_content_lines.append(text)

            if re.match(r'^Пример\s+\d+', text, re.IGNORECASE):
                ex_lines = [text]
                j = i + 1
                while j < len(paragraphs) and j < i + 20:
                    nt = extract_paragraph_full_text(paragraphs[j]).strip()
                    if nt and not re.match(r'^Пример\s+\d+', nt, re.IGNORECASE):
                        ex_lines.append(nt)
                        j += 1
                    else:
                        break
                examples.append("\n".join(ex_lines)[:800])

        i += 1

    flush_section()

    for section in sections:
        if "самостоятельн" in section["title"].lower():
            lines = section["content"].split("\n")
            current_variant = ""
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                if re.match(r'^Вариант\s+\d+', line, re.IGNORECASE):
                    current_variant = line
                elif current_variant and len(line) > 20 and not any(
                        x in line.lower() for x in ["ответ", "решение", "список литературы"]
                ):
                    tasks.append(f"[{section['title']} — {current_variant}]\n{line[:500]}")

    tasks = list(dict.fromkeys(tasks))

    print(f"\nРезультаты парсинга:")
    print(f"  Разделов: {len(sections)}, Примеров: {len(examples)}, Задач: {len(tasks)}")
    if tasks:
        print("  Первые 3 задачи:")
        for t in tasks[:3]:
            print(f"    • {t[:80].replace(chr(10), ' ')}")

    return {
        "full_text": TEXT,
        "sections": sections,
        "tasks": tasks,
        "examples": examples,
        "answers": answers
    }


TEXT = load_docx_text()
BOOK = parse_book_from_docx()


def find_relevant_context(query, max_chars=MAX_CONTEXT_CHARS):
    if not BOOK["sections"]:
        return BOOK["full_text"][:max_chars] if BOOK["full_text"] else "Учебник не загружен"

    query_lower = query.lower()
    query_words = set(query_lower.split()) - {
        'это', 'как', 'что', 'для', 'на', 'в', 'с', 'по', 'из', 'за',
        'у', 'о', 'к', 'и', 'а', 'но', 'да', 'нет', 'или', 'же', 'мне',
        'мой', 'моя', 'дай', 'дайте', 'пожалуйста'
    }

    scored = []
    for section in BOOK["sections"]:
        combined = f"{section['title']}\n{section['content']}".lower()
        score = sum(1 for w in query_words if w in combined)
        if query_lower[:20] in combined:
            score += 5
        if score > 0:
            scored.append((score, section))

    scored.sort(key=lambda x: x[0], reverse=True)

    context_parts = []
    total = 0
    for _, section in scored[:MAX_SEARCH_BLOCKS]:
        block = f"{section['title']}\n\n{section['content'][:2000]}"
        if total + len(block) > max_chars:
            break
        context_parts.append(block)
        total += len(block)

    if context_parts:
        return "\n\n---\n\n".join(context_parts)

    if BOOK["sections"]:
        s = BOOK["sections"][0]
        return f"{s['title']}\n\n{s['content'][:max_chars]}"

    return BOOK["full_text"][:max_chars] if BOOK["full_text"] else "Учебник не загружен"


def clean_latex(text):
    """Преобразует LaTeX в читаемый текст для Telegram с юникодными индексами"""
    text = re.sub(r'\$\$[^\$]+\$\$', '', text)
    text = re.sub(r'\$[^\$]+\$', '', text)

    greek = {
        r'\\alpha': 'α', r'\\beta': 'β', r'\\gamma': 'γ', r'\\delta': 'δ',
        r'\\epsilon': 'ε', r'\\varepsilon': 'ε', r'\\zeta': 'ζ', r'\\eta': 'η',
        r'\\theta': 'θ', r'\\iota': 'ι', r'\\kappa': 'κ', r'\\lambda': 'λ',
        r'\\mu': 'μ', r'\\nu': 'ν', r'\\xi': 'ξ', r'\\pi': 'π', r'\\rho': 'ρ',
        r'\\sigma': 'σ', r'\\tau': 'τ', r'\\upsilon': 'υ', r'\\phi': 'φ',
        r'\\chi': 'χ', r'\\psi': 'ψ', r'\\omega': 'ω',
        r'\\Gamma': 'Γ', r'\\Delta': 'Δ', r'\\Theta': 'Θ', r'\\Lambda': 'Λ',
        r'\\Xi': 'Ξ', r'\\Pi': 'Π', r'\\Sigma': 'Σ', r'\\Phi': 'Φ',
        r'\\Psi': 'Ψ', r'\\Omega': 'Ω'
    }
    for latex, symbol in greek.items():
        text = text.replace(latex, symbol)

    sub_letters = {
        'a': 'ₐ', 'b': '♭', 'c': 'c', 'd': 'd', 'e': 'ₑ', 'f': 'f', 'g': 'g',
        'h': 'ₕ', 'i': 'ᵢ', 'j': 'ⱼ', 'k': 'ₖ', 'l': 'ₗ', 'm': 'ₘ', 'n': 'ₙ',
        'o': 'ₒ', 'p': 'ₚ', 'q': 'q', 'r': 'ᵣ', 's': 'ₛ', 't': 'ₜ', 'u': 'ᵤ',
        'v': 'ᵥ', 'w': 'w', 'x': 'ₓ', 'y': 'y', 'z': 'z'
    }

    sub_digits = {
        '0': '₀', '1': '₁', '2': '₂', '3': '₃', '4': '₄',
        '5': '₅', '6': '₆', '7': '₇', '8': '₈', '9': '₉'
    }

    sup_letters = {
        'a': 'ᵃ', 'b': 'ᵇ', 'c': 'ᶜ', 'd': 'ᵈ', 'e': 'ᵉ', 'f': 'ᶠ', 'g': 'ᵍ',
        'h': 'ʰ', 'i': 'ⁱ', 'j': 'ʲ', 'k': 'ᵏ', 'l': 'ˡ', 'm': 'ᵐ', 'n': 'ⁿ',
        'o': 'ᵒ', 'p': 'ᵖ', 'q': 'ʳ', 'r': 'ʳ', 's': 'ˢ', 't': 'ᵗ', 'u': 'ᵘ',
        'v': 'ᵛ', 'w': 'ʷ', 'x': 'ˣ', 'y': 'ʸ', 'z': 'ᶻ'
    }

    sup_digits = {
        '0': '⁰', '1': '¹', '2': '²', '3': '³', '4': '⁴',
        '5': '⁵', '6': '⁶', '7': '⁷', '8': '⁸', '9': '⁹'
    }

    def convert_subscript(match):
        content = match.group(1) if match.group(1) else match.group(2)
        result = []
        for ch in content:
            if ch in sub_digits:
                result.append(sub_digits[ch])
            elif ch in sub_letters:
                result.append(sub_letters[ch])
            else:
                result.append(ch)
        return ''.join(result)

    def convert_superscript(match):
        content = match.group(1) if match.group(1) else match.group(2)
        result = []
        for ch in content:
            if ch in sup_digits:
                result.append(sup_digits[ch])
            elif ch in sup_letters:
                result.append(sup_letters[ch])
            else:
                result.append(ch)
        return ''.join(result)

    text = re.sub(r'\_\{([^{}]+)\}', convert_subscript, text)
    text = re.sub(r'\_([a-zA-Z0-9])', lambda m: convert_subscript(m), text)
    text = re.sub(r'\^\{([^{}]+)\}', convert_superscript, text)
    text = re.sub(r'\^([a-zA-Z0-9])', lambda m: convert_superscript(m), text)

    text = re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'\1/\2', text)

    replacements = {
        r'\\cdot': '·', r'\\times': '×', r'\\to': '→', r'\\infty': '∞',
        r'\\sum': '∑', r'\\prod': '∏', r'\\int': '∫', r'\\partial': '∂',
        r'\\sqrt': '√', r'\\neq': '≠', r'\\leq': '≤', r'\\geq': '≥',
        r'\\approx': '≈', r'\\ldots': '…', r'\\cdots': '⋯', r'\\vdots': '⋮', r'\\ddots': '⋱'
    }
    for latex, symbol in replacements.items():
        text = text.replace(latex, symbol)

    text = re.sub(r'\\[a-zA-Z]+(\{[^}]*\})*', '', text)
    text = re.sub(r'\{|\}', '', text)
    text = re.sub(r' +', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)

    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)

    return text.strip()


def clean_response(text):
    """Очищает ответ от LaTeX и форматирует для Telegram"""
    text = clean_latex(text)

    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', text)
    text = re.sub(r'__([^_]+)__', r'<b>\1</b>', text)
    text = re.sub(r'_([^_]+)_', r'<i>\1</i>', text)
    text = re.sub(r'^#{1,3}\s+(.+)', r'<b>\1</b>', text, flags=re.MULTILINE)
    text = re.sub(r'`([^`]+)`', r'<code>\1</code>', text)
    text = re.sub(r'```[\w]*\n?(.*?)```', r'<pre>\1</pre>', text, flags=re.DOTALL)

    return text.strip()


SYSTEM_PROMPT = """Ты преподаватель по тензорному исчислению. Работаешь по учебному пособию.

КЛЮЧЕВЫЕ ТРЕБОВАНИЯ К ФОРМАТИРОВАНИЮ:
1. Используй обычный текст БЕЗ LaTeX (без знаков $ и \\)
2. Индексы пиши через подчеркивание: a_i, b_j^k
3. Греческие буквы пиши как символы юникод: α, β, γ
4. Дроби: a/b
5. Умножение: точка или пробел

СТИЛЬ ОТВЕТА:
- Кратко и по делу (максимум 4-5 коротких абзацев)
- Для определений - 2-3 предложения
- Для задач - пошагово, каждый шаг с новой строки
- Не перечисляй всю теорию подряд

Основавайся на контексте из учебника. Если темы нет - скажи коротко."""


def ask(messages, max_tokens=500, retry_count=3):
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://t.me/tensor_bot",
        "X-OpenRouter-Title": "Tensor Assistant"
    }
    payload = {
        "model": MODEL,
        "messages": messages,
        "max_tokens": max_tokens,
        "temperature": 0.3
    }

    for attempt in range(retry_count):
        try:
            response = requests.post(URL, headers=headers, json=payload, timeout=90, verify=False)
            if response.status_code == 200:
                result = response.json()
                if "choices" in result and result["choices"]:
                    content = result["choices"][0].get("message", {}).get("content", "")
                    return content if content else "Модель вернула пустой ответ"
                return "Неожиданный формат ответа"
            else:
                if attempt < retry_count - 1:
                    time.sleep(2)
                    continue
                return f"Ошибка API ({response.status_code}): {response.text[:200]}"
        except Exception as e:
            if attempt < retry_count - 1:
                time.sleep(2)
                continue
            return f"Ошибка запроса: {str(e)[:100]}"

    return "Не удалось получить ответ."


def is_request_task(text):
    t = text.lower()
    return any(x in t for x in [
        "дай задачу", "пришли задачу", "хочу задачу", "ещё задачу", "еще задачу",
        "задачу пожалуйста", "дайте задачу", "следующую задачу", "новую задачу",
        "дай другую задачу", "другую задачу", "задачу", "/task"
    ])


def is_request_answer(text):
    t = text.lower()
    return any(x in t for x in [
        "покажи ответ", "дай ответ", "какой ответ", "покажи решение", "ответ"
    ])


def is_request_hint(text):
    t = text.lower()
    return any(x in t for x in [
        "подскажи", "подсказка", "дай подсказку", "помоги", "не понимаю"
    ])


def process_text(user_id, text):
    user_id_str = str(user_id)
    conversation = user_conversations.get(user_id_str, [])

    # Проверяем запрос на задачу
    if is_request_task(text):
        task = get_random_task_from_json()
        if not task:
            return "Задачи не найдены. Проверьте файл tasks.json"

        # Сохраняем текущую задачу пользователя
        user_current_task[user_id_str] = task["id"]
        save_memory()

        conversation.append({"role": "user", "content": text})
        conversation.append({"role": "assistant", "content": f"[ЗАДАЧА_ID:{task['id']}]{task['question']}"})
        if len(conversation) > 20:
            conversation = conversation[-20:]
        user_conversations[user_id_str] = conversation
        save_memory()

        return format_task_from_json(task)

    # Проверяем запрос на подсказку
    if is_request_hint(text):
        current_task_id = user_current_task.get(user_id_str)
        if current_task_id:
            task = get_task_by_id(current_task_id)
            if task:
                return format_hint_from_json(task)
            else:
                return "Задача не найдена. Попроси новую задачу командой /task"
        else:
            return "У тебя нет активной задачи. Напиши /task чтобы получить задачу"

    # Проверяем запрос на ответ
    if is_request_answer(text):
        current_task_id = user_current_task.get(user_id_str)
        if current_task_id:
            task = get_task_by_id(current_task_id)
            if task:
                # Если пользователь просто просит ответ
                if any(x in text.lower() for x in ["покажи ответ", "дай ответ", "какой ответ", "покажи решение"]):
                    return format_answer_from_json(task)

                # Если пользователь прислал своё решение - проверяем через нейросеть
                context = find_relevant_context(task['question'])

                user_prompt = f"""Задача из учебника:
{task['question']}

Эталонный ответ:
{task['answer']}

Ответ студента:
{text}

Пожалуйста, проверь ответ студента. Сравни с эталонным ответом.
Если ответ верный - похвали и скажи, что правильно.
Если есть ошибки - укажи их и дай подсказку как исправить.
Будь доброжелательным и конструктивным.

Ответ должен быть кратким (2-4 предложения)."""

                messages = [
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {"role": "user", "content": user_prompt}
                ]

                final = clean_response(ask(messages, max_tokens=500))

                conversation.append({"role": "user", "content": text})
                conversation.append({"role": "assistant", "content": final})
                if len(conversation) > 20:
                    conversation = conversation[-20:]
                user_conversations[user_id_str] = conversation
                save_memory()

                return final
            else:
                return "Задача не найдена. Попроси новую задачу командой /task"
        else:
            return "У тебя нет активной задачи. Напиши /task чтобы получить задачу"

    # Обычный вопрос - используем контекст из учебника
    context = find_relevant_context(text)

    messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        *conversation[-6:],
        {
            "role": "user",
            "content": f"Контекст из учебника:\n{context}\n\n---\n\nВопрос: {text}\n\nОтветь кратко, структурируя ответ. Каждое новое предложение или пункт с новой строки. Не используй LaTeX."
        }
    ]

    final = clean_response(ask(messages, max_tokens=500))

    conversation.append({"role": "user", "content": text})
    conversation.append({"role": "assistant", "content": final})
    if len(conversation) > 20:
        conversation = conversation[-20:]
    user_conversations[user_id_str] = conversation
    save_memory()

    return final


def process_image(user_id, image_bytes):
    try:
        img = Image.open(io.BytesIO(image_bytes))
        if img.mode != 'RGB':
            img = img.convert('RGB')
        if img.width > 1024 or img.height > 1024:
            ratio = min(1024 / img.width, 1024 / img.height)
            img = img.resize((int(img.width * ratio), int(img.height * ratio)), Image.Resampling.LANCZOS)
        buf = io.BytesIO()
        img.save(buf, format='JPEG', quality=85)
        buf.seek(0)
        b64 = base64.b64encode(buf.read()).decode()
        context = find_relevant_context("задача тензор", max_chars=2000)
        messages = [
            {"role": "system", "content": SYSTEM_PROMPT},
            {
                "role": "user",
                "content": [
                    {"type": "text",
                     "text": f"Реши задачу с изображения. Контекст:\n{context}\n\nНе используй LaTeX. Каждый шаг с новой строки."},
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}}
                ]
            }
        ]
        return clean_response(ask(messages, max_tokens=800))
    except Exception as e:
        return f"Ошибка при обработке изображения: {str(e)[:100]}"


def send_reply(message, text):
    if not text:
        text = "Извините, не могу ответить на этот запрос."
    chunks = [text[i:i + 4000] for i in range(0, len(text), 4000)]
    for chunk in chunks:
        try:
            bot.reply_to(message, chunk, parse_mode='HTML')
        except Exception:
            plain = re.sub(r'<[^>]+>', '', chunk)
            bot.reply_to(message, plain)


@bot.message_handler(commands=['start'])
def start(message):
    bot.reply_to(
        message,
        f"<b>Привет! Я твой персональный преподаватель по тензорам!</b>\n\n"
        f"<b>Статистика:</b>\n"
        f"• Задач в базе: {len(TASKS)}\n\n"
        f"<b>Команды:</b>\n"
        f"/task — получить случайную задачу\n"
        f"/reset — сбросить контекст\n\n"
        f"После получения задачи:\n"
        f"- Пришли своё решение — я проверю!\n"
        f"Ты можешь также запросить ответ или подсказку :)\n",
        parse_mode='HTML'
    )



@bot.message_handler(commands=['reset'])
def reset(message):
    uid = str(message.from_user.id)
    if uid in user_conversations:
        del user_conversations[uid]
    if uid in user_current_task:
        del user_current_task[uid]
    save_memory()
    bot.reply_to(message, "Контекст и текущая задача сброшены.")


@bot.message_handler(commands=['task'])
def send_task(message):
    bot.send_chat_action(message.chat.id, 'typing')

    task = get_random_task_from_json()
    if task:
        send_reply(message, format_task_from_json(task))

        uid = str(message.from_user.id)
        user_current_task[uid] = task["id"]
        save_memory()

        conv = user_conversations.get(uid, [])
        conv.append({"role": "user", "content": "/task"})
        conv.append({"role": "assistant", "content": f"[ЗАДАЧА_ID:{task['id']}]{task['question']}"})
        if len(conv) > 20:
            conv = conv[-20:]
        user_conversations[uid] = conv
        save_memory()
    else:
        bot.reply_to(message, "База задач не загружена. Проверьте файл tasks.json")



@bot.message_handler(commands=['sections'])
def show_sections(message):
    if not BOOK["sections"]:
        bot.reply_to(message, "Разделы не найдены. Проверьте файл text.docx")
        return
    lines = ["<b>Разделы учебника:</b>\n"]
    for i, s in enumerate(BOOK["sections"][:20], 1):
        lines.append(f"{i}. {s['title'][:60]}")
    if len(BOOK["sections"]) > 20:
        lines.append(f"\n...и ещё {len(BOOK['sections']) - 20}")
    bot.reply_to(message, "\n".join(lines), parse_mode='HTML')


@bot.message_handler(commands=['debug'])
def debug(message):
    info = (
        f"<b>Отладка:</b>\n\n"
        f"Текст загружен: {'да' if TEXT else 'нет'}\n"
        f"Разделов: {len(BOOK['sections'])}\n"
        f"Задач в JSON: {len(TASKS)}\n"
        f"Задач в DOCX: {len(BOOK['tasks'])}\n"
    )
    bot.reply_to(message, info, parse_mode='HTML')


@bot.message_handler(content_types=['photo'])
def handle_photo(message):
    try:
        bot.send_chat_action(message.chat.id, 'typing')
        bot.reply_to(message, "Анализирую изображение...")
        file_info = bot.get_file(message.photo[-1].file_id)
        file_bytes = bot.download_file(file_info.file_path)
        answer = process_image(message.from_user.id, file_bytes)
        send_reply(message, answer)
    except Exception as e:
        bot.reply_to(message, f"Ошибка: {str(e)[:100]}")
        logger.error(f"Photo error: {e}")


@bot.message_handler(func=lambda m: True)
def handle_text(message):
    if message.text.startswith("/"):
        return
    bot.send_chat_action(message.chat.id, 'typing')
    answer = process_text(message.from_user.id, message.text)
    send_reply(message, answer)


if __name__ == "__main__":
    load_memory()

    print("БОТ ЗАПУЩЕН")
    print(f"Модель: {MODEL}")
    print(f"Задач в JSON: {len(TASKS)}")
    print(f"Учебник: {'загружен' if TEXT else 'НЕ НАЙДЕН'}")
    print(f"Разделов: {len(BOOK['sections'])}")

    try:
        bot.infinity_polling(timeout=60)
    except KeyboardInterrupt:
        print("\nБот остановлен")
    except Exception as e:
        print(f"Ошибка: {e}")